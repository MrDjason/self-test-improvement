# Day 26 操作Word和PPT
# ==================== 导入模块 ====================
from docx import Document
from docx.shared import Cm, Pt
# 从docx.shared导入Cm（厘米单位）和Pt（磅单位）
from docx.document import Document as Doc

import os
# ==================== 主 程 序 ====================
script_dir = os.path.dirname(os.path.abspath(__file__))
file_path = os.path.join(script_dir, 'res')

document = Document()
# 创建代表Word文档的Doc对象

document.add_heading('快快乐乐学python', 0)
# 添加0级大标题

p = document.add_paragraph('Python是一门非常流行的编程语言，它')
# 添加段落

run = p.add_run('简单易学')
# 向段落p中添加一个文本片段"简单"，用run对象表示

run.bold = True        # 加粗
run.font.size = Pt(18) # 18磅

p.add_run('而且')
# 继续向段落p添加文本"而且"
run = p.add_run('优雅')
# 向段落p中添加另一个文本片段"优雅"

run.font.size = Pt(18)
run.underline = True
# 设置该文本片段的字体大小为18磅，且有下划线

p.add_run('。')
# 继续向段落p添加文本"，完成整个段落

document.add_heading('Heading, level 1', level=1)
# 添加一级标题

document.add_paragraph('Intense quote', style = 'Intense Quote')
# 添加带样式的段落

document.add_paragraph('first item in unordered list', style='List Bullet')
# 添加无序列表

document.add_paragraph('second item in ordered list', style='List Bullet')

document.add_paragraph('first item in ordered list', style='List Number')
# 添加有序列表

document.add_paragraph('second item in ordered list', style='List Number')
document.add_picture(file_path +'/test1.jpg', width=Cm(5.2)) 
# 添加图片

document.add_section()
# 添加分节符

records = (
    ('李明', '男', '1995-5-5'),
    ('孙丽', '女', '1992-2-2')
)

# 添加表格
table = document.add_table(rows=1,cols=3)
table.style = 'Dark List'
hdr_cells = table.rows[0].cells # 获取表格第一行（表头）的单元格样式
hdr_cells[0].text = '姓名'
hdr_cells[1].text = '性别'
hdr_cells[2].text = '出生日期'
for name, sex, birthday in records:
     # 为表格添加一行，并获取该行的单元格列表
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = sex
    row_cells[2].text = birthday

document.add_page_break()
# 添加分页符


document.save(file_path + '/demo.docx')
# 保存文档

from docx import Document
from docx.document import Document as Doc

doc = Document(file_path + '/离职证明.docx')
for no, p in enumerate(doc.paragraphs):
    print(no, p.text)
    # # no：段落的索引（从0开始，用于标识段落位置）
    # p：每个段落对象（docx.text.paragraph.Paragraph），代表文档中的一个段落

# 我们可以把上面的离职证明制作成一个模板文件，把姓名、身份证号、入职和离职日期等信息用占位符代替
# 这样通过对占位符的替换，就可以根据实际需要写入对应的信息